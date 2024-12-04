import spacy
from pathlib import Path
from docx import Document
from fpdf import FPDF


def load_language_model(language):
    """Load the appropriate spaCy language model."""
    if language == "english":
        return spacy.load("en_core_web_sm")
    elif language == "hebrew":
        return spacy.load("he_core_news_sm")
    else:
        raise ValueError("Unsupported language. Choose either 'english' or 'hebrew'.")


def read_file(input_path):
    """Read text from txt or docx file."""
    input_path = Path(input_path)
    if input_path.suffix == ".txt":
        with open(input_path, "r", encoding="utf-8") as file:
            return file.read()
    elif input_path.suffix == ".docx":
        doc = Document(input_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type. Supported types are: .txt, .docx")


def anonymize_text(text, nlp):
    """Anonymize text using spaCy."""
    doc = nlp(text)
    replacements = []
    placeholder_mapping = {
        "PERSON": "אדם" if nlp.lang == "he" else "person",
        "DATE": "תאריך" if nlp.lang == "he" else "date",
        "GPE": "מיקום" if nlp.lang == "he" else "location",
    }
    placeholder_counter = {}

    for ent in doc.ents:
        if ent.label_ in placeholder_mapping:
            placeholder_type = placeholder_mapping[ent.label_]
            if ent.text not in placeholder_counter:
                placeholder_counter[ent.text] = f"{placeholder_type}_{len(placeholder_counter) + 1}"
            replacements.append((ent.start_char, ent.end_char, placeholder_counter[ent.text]))

    for start, end, placeholder in sorted(replacements, key=lambda x: x[0], reverse=True):
        text = text[:start] + placeholder + text[end:]

    return text


def write_file(output_path, text, format):
    """Write anonymized text to the desired output format."""
    output_path = Path(output_path)
    if format == "txt":
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(text)
    elif format == "docx":
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(output_path)
    elif format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.split("\n"):
            pdf.multi_cell(0, 10, line)
        pdf.output(output_path)
    else:
        raise ValueError("Unsupported output format. Supported formats are: txt, docx, pdf.")


def main(input_path, output_path, language, output_format):
    """Main function to handle the anonymization process."""
    # Load the appropriate language model
    nlp = load_language_model(language)

    # Read the input file
    text = read_file(input_path)

    # Anonymize the text
    anonymized_text = anonymize_text(text, nlp)

    # Write the output to the desired location and format
    write_file(output_path, anonymized_text, output_format)

    print(f"Anonymized file created at: {output_path}")


# Example usage
if __name__ == "__main__":
    # Input params
    input_path = "path/to/your/input/file.txt"  # Replace with the actual path
    output_path = "path/to/your/output/file.pdf"  # Replace with the desired output path
    language = "english"  # or "hebrew"
    output_format = "pdf"  # txt, docx, or pdf

    main(input_path, output_path, language, output_format)
